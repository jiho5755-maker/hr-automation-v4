"""
급여관리 자동화 - 메인 애플리케이션
4대보험, 소득세 자동 계산 및 급여명세서 생성
"""

import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
from datetime import datetime, timedelta
from io import BytesIO
import sys
from pathlib import Path

# shared 모듈 import
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))
from shared.database import get_all_employees, get_employee_by_id, get_company_profile
from shared.design import apply_design
from shared.utils import show_success

# 로컬 모듈 import
import constants as C
from calculator import (
    PayrollCalculator,
    AnnualLeaveCalculator,
    calculate_hourly_wage,
    calculate_overtime_pay,
    calculate_ot_hours_from_pay,
    format_payslip,
    validate_working_hours,
    validate_minimum_wage
)
from database import (
    init_payroll_tables,
    add_payroll_setting,
    get_payroll_setting,
    update_payroll_setting,
    get_all_payroll_settings,
    add_payroll_history,
    get_payroll_history,
    get_employee_payroll_history,
    get_monthly_payroll_summary,
    update_paid_status,
    add_overtime_log,
    get_monthly_overtime,
    add_annual_leave,
    get_annual_leave,
    add_annual_leave_usage,
    get_annual_leave_usage,
    init_annual_leave_if_not_exists
)

# ============================================================
# 페이지 설정
# ============================================================

st.set_page_config(
    page_title=C.APP_CONFIG["제목"],
    page_icon=C.APP_CONFIG["아이콘"],
    layout=C.APP_CONFIG["레이아웃"],
    initial_sidebar_state="expanded"
)

# 모던 그린 미니멀 디자인 적용
apply_design()

# ============================================================
# 세션 상태 초기화
# ============================================================

def init_session_state():
    """세션 상태 초기화"""
    if 'payroll_calculator' not in st.session_state:
        # 회사 정보에서 직원 수 가져오기
        company = get_company_profile()
        employee_count = company['employee_count'] if company else 1
        st.session_state.payroll_calculator = PayrollCalculator(employee_count)
    
    if 'current_year_month' not in st.session_state:
        st.session_state.current_year_month = C.get_current_year_month()

init_session_state()

# ============================================================
# 메인 페이지
# ============================================================

st.markdown('<div class="main-title">💰 급여관리 자동화</div>', unsafe_allow_html=True)

st.markdown(f"""
**{C.APP_CONFIG["설명"]}**

✨ **주요 기능**
- 📊 4대보험 자동 계산 (국민연금, 건강보험, 고용보험, 산재보험)
- 💵 소득세/지방소득세 자동 계산
- 📄 급여명세서 자동 생성 및 PDF 출력
- 📈 급여대장 엑셀 다운로드
- 🎯 시간외 수당 계산 (연장/야간/휴일)
- 📅 연차수당 자동 계산
- 💼 포괄임금제 지원 (DC형 퇴직연금 적용)
""")

st.divider()

# ============================================================
# 메뉴 선택
# ============================================================

# 홈 버튼
st.sidebar.markdown("### 🏠 메뉴")
if st.sidebar.button("🏠 통합 대시보드로 이동", use_container_width=True, key="home_dashboard"):
    st.sidebar.info("🌐 브라우저에서 http://localhost:8000 로 접속하세요")

st.sidebar.divider()

# 급여 설정 안내
st.sidebar.info("""
⚙️ **급여 정보 설정**

통합 대시보드에서:
📍 http://localhost:8000
→ 💰 급여 정보 관리
""")

st.sidebar.divider()

menu = st.sidebar.selectbox(
    "📌 기능 메뉴",
    [
        "🏠 대시보드",
        "💰 월별 급여 계산",
        "📊 급여대장",
        "📄 급여명세서 출력",
        "⏰ 시간외 수당",
        "📅 연차 관리"
    ]
)

# 귀속 년월 선택
st.sidebar.divider()
st.sidebar.markdown("### 📅 귀속 년월")
year_month = st.sidebar.text_input(
    "년월 (YYYY-MM)",
    value=st.session_state.current_year_month,
    help="급여 계산 기준 년월"
)

# ============================================================
# 대시보드
# ============================================================

if menu == "🏠 대시보드":
    st.subheader("📊 급여 현황")
    
    # 통계
    col1, col2, col3, col4 = st.columns(4)
    
    # 직원 수
    employees = get_all_employees(active_only=True)
    emp_count = len(employees)
    
    # 급여 설정된 직원 수
    payroll_settings = get_all_payroll_settings()
    payroll_count = len(payroll_settings)
    
    # 이번 달 급여 계산 완료 직원 수
    monthly_payroll = get_monthly_payroll_summary(year_month)
    calculated_count = len(monthly_payroll)
    
    # 총 급여 지급액
    total_payment = sum([p['net_pay'] for p in monthly_payroll])
    
    with col1:
        st.metric("👥 전체 직원", f"{emp_count}명")
    with col2:
        st.metric("⚙️ 급여 설정", f"{payroll_count}명")
    with col3:
        st.metric("✅ 계산 완료", f"{calculated_count}명")
    with col4:
        st.metric("💰 총 지급액", C.format_currency(total_payment))
    
    st.divider()
    
    # 이번 달 급여 요약
    if monthly_payroll:
        st.subheader(f"📋 {year_month} 급여 요약")
        
        df = pd.DataFrame(monthly_payroll)
        df['base_salary'] = df['base_salary'].apply(C.format_currency)
        df['total_allowance'] = df['total_allowance'].apply(C.format_currency)
        df['total_deduction'] = df['total_deduction'].apply(C.format_currency)
        df['net_pay'] = df['net_pay'].apply(C.format_currency)
        
        df = df.rename(columns={
            'name': '성명',
            'department': '부서',
            'position': '직급',
            'base_salary': '기본급',
            'total_allowance': '수당',
            'total_deduction': '공제',
            'net_pay': '실수령액',
            'paid_status': '지급상태'
        })
        
        st.dataframe(
            df[['성명', '부서', '직급', '기본급', '수당', '공제', '실수령액', '지급상태']],
            use_container_width=True,
            hide_index=True
        )
    else:
        st.info(f"💡 {year_month} 급여가 아직 계산되지 않았습니다. '💰 월별 급여 계산' 메뉴에서 계산하세요.")
    
    # 미설정 직원 알림
    st.divider()
    unset_employees = [emp for emp in employees if not any(ps['emp_id'] == emp['emp_id'] for ps in payroll_settings)]
    
    if unset_employees:
        st.warning(f"⚠️ **급여 미설정 직원**: {len(unset_employees)}명")
        for emp in unset_employees:
            st.write(f"- {emp['name']} ({emp['department']} / {emp['position']})")
        st.info("""
        💡 **급여 정보를 입력하려면?**
        
        👉 통합 대시보드(포트 8000)에서 설정하세요!
        
        📍 http://localhost:8000 → 💰 급여 정보 관리
        """)
    else:
        st.success("✅ 모든 직원의 급여가 설정되었습니다!")

# ============================================================
# 월별 급여 계산
# ============================================================

elif menu == "💰 월별 급여 계산":
    st.subheader(f"💰 {year_month} 급여 계산")
    
    # 직원 선택
    employees = get_all_employees(active_only=True)
    employee_options = {f"{emp['name']} ({emp['department']})": emp for emp in employees}
    
    selected = st.selectbox("👤 직원 선택", list(employee_options.keys()))
    
    if selected:
        employee = employee_options[selected]
        emp_id = employee['emp_id']
        
        # 급여 설정 확인
        setting = get_payroll_setting(emp_id)
        
        if not setting:
            st.warning(f"""
            ⚠️ **{employee['name']}님의 급여 설정이 없습니다!**
            
            👉 통합 대시보드에서 먼저 설정하세요.
            
            📍 http://localhost:8000
            → 💰 급여 정보 관리
            → {employee['name']} 선택 → 저장
            """)
            st.stop()
        else:
            # 일할계산 옵션
            st.markdown("#### 📅 일할계산")
            use_prorated = st.checkbox(
                "일할계산 적용",
                help="월 중 입/퇴사자나 휴직자 등의 일할 계산"
            )
            
            work_days = None
            month_days = None
            
            if use_prorated:
                col1, col2 = st.columns(2)
                with col1:
                    work_days = st.number_input(
                        "실 근무일수",
                        min_value=1,
                        max_value=31,
                        value=15,
                        help="해당 월의 실제 근무일수"
                    )
                with col2:
                    month_days = st.number_input(
                        "월 총 일수",
                        min_value=28,
                        max_value=31,
                        value=31,
                        help="해당 월의 전체 일수"
                    )
                
                st.info(f"💡 일할 계산: {work_days}/{month_days}일 = {work_days/month_days*100:.1f}%")
            
            st.divider()
            
            # 급여 계산
            calc_result = st.session_state.payroll_calculator.calculate_all(
                base_salary=setting['base_salary'],
                allowances=setting['allowances'],
                tax_free_items=setting['tax_free_items'],
                apply_pension=setting.get('apply_pension', True),
                apply_health=setting.get('apply_health', True),
                apply_longterm=setting.get('apply_longterm', True),
                apply_employment=setting.get('apply_employment', True),
                fixed_ot_amount=setting.get('fixed_ot_amount', 0),
                work_days=work_days,
                month_days=month_days
            )
            
            st.divider()
            
            # 지급 내역
            st.markdown("### 💵 지급 내역")
            
            # 기본급
            st.markdown("#### 💰 기본급")
            st.metric("기본급", C.format_currency(calc_result['기본급']))
            
            # 수당 내역 상세 표시 (1월 급여대장 형식)
            st.markdown("#### 🎁 수당 내역")
            if calc_result['수당']:
                # 주요 수당 분류
                meal_allowance = calc_result['수당'].get('식대', 0)
                transport_allowance = calc_result['수당'].get('교통비', 0)
                overtime_total = (calc_result['수당'].get('연장근로수당', 0) + 
                                 calc_result['수당'].get('야간근로수당', 0) + 
                                 calc_result['수당'].get('휴일근로수당', 0))
                other_total = sum([v for k, v in calc_result['수당'].items() 
                                 if k not in ['식대', '교통비', '연장근로수당', '야간근로수당', '휴일근로수당']])
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("식대", C.format_currency(meal_allowance), help="🔵 비과세 (월 20만원 한도)")
                with col2:
                    st.metric("교통비", C.format_currency(transport_allowance), help="🔵 비과세 (월 20만원 한도)")
                with col3:
                    st.metric("연장/야간/휴일수당", C.format_currency(overtime_total), help="🟢 과세")
                with col4:
                    st.metric("기타수당", C.format_currency(other_total), help="🟢 과세")
                
                st.divider()
                col_sum1, col_sum2, col_sum3 = st.columns(3)
                with col_sum1:
                    st.metric("**수당 합계**", C.format_currency(calc_result['총수당']))
                with col_sum2:
                    st.metric("**비과세 합계**", C.format_currency(calc_result.get('총비과세', 0)))
                with col_sum3:
                    st.metric("**총 지급액**", C.format_currency(calc_result['총지급액']))
            
            # 공제 내역
            st.divider()
            st.markdown("### 🧾 공제 내역")
            
            # 4대 사회보험
            st.markdown("#### 🏥 4대 사회보험")
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("국민연금", C.format_currency(calc_result['국민연금']))
                st.caption("근로자 4.75%")
            with col2:
                st.metric("건강보험", C.format_currency(calc_result['건강보험']))
                st.caption("근로자 3.60%")
            with col3:
                st.metric("장기요양", C.format_currency(calc_result['장기요양']))
                st.caption("건강보험료의 13.14%")
            with col4:
                st.metric("고용보험", C.format_currency(calc_result['고용보험']))
                st.caption("근로자 0.9%")
            
            # 세금
            st.markdown("#### 💵 세금")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("소득세", C.format_currency(calc_result['소득세']))
                st.caption("간이세액표 기준")
            with col2:
                st.metric("지방소득세", C.format_currency(calc_result['지방세']))
                st.caption("소득세의 10%")
            with col3:
                st.metric("**총 공제액**", C.format_currency(calc_result['총공제']))
                st.caption("4대보험 + 세금")
            
            # 실수령액
            st.divider()
            st.markdown("### 💰 실수령액")
            st.markdown(f"""
            <div class="metric-card">
                <h2>{C.format_currency(calc_result['실수령액'])}</h2>
                <p>실제 지급 금액</p>
            </div>
            """, unsafe_allow_html=True)
            
            # 저장 버튼
            st.divider()
            if st.button("💾 급여 이력에 저장", use_container_width=True, type="primary"):
                # 데이터베이스 저장용 형식으로 변환
                payroll_history_data = {
                    '지급내역': {
                        '기본급': calc_result['기본급'],
                        '수당합계': calc_result['총수당'],
                        '과세대상액': calc_result['과세급여']
                    },
                    '공제내역': {
                        '국민연금': calc_result['국민연금'],
                        '건강보험': calc_result['건강보험'],
                        '장기요양': calc_result['장기요양'],
                        '고용보험': calc_result['고용보험'],
                        '소득세': calc_result['소득세'],
                        '지방소득세': calc_result['지방세'],
                        '공제합계': calc_result['총공제']
                    },
                    '실수령액': calc_result['실수령액'],
                    '수당상세': calc_result['수당']
                }
                
                if add_payroll_history(emp_id, payroll_history_data, year_month):
                    show_success("급여 이력이 저장되었습니다!")
                else:
                    st.error("❌ 급여 이력 저장 실패")
            
            # 상세 정보 (접기)
            with st.expander("📊 상세 정보 보기"):
                st.json(calc_result, expanded=False)

# ============================================================
# 급여대장
# ============================================================

elif menu == "📊 급여대장":
    st.subheader(f"📊 {year_month} 급여대장")
    
    monthly_payroll = get_monthly_payroll_summary(year_month)
    
    if not monthly_payroll:
        st.info(f"💡 {year_month} 급여 데이터가 없습니다.")
    else:
        # DataFrame 생성
        df = pd.DataFrame(monthly_payroll)
        
        # 통계
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("총 인원", f"{len(df)}명")
        with col2:
            st.metric("총 지급액", C.format_currency(df['base_salary'].sum() + df['total_allowance'].sum()))
        with col3:
            st.metric("총 공제액", C.format_currency(df['total_deduction'].sum()))
        with col4:
            st.metric("실수령액 합계", C.format_currency(df['net_pay'].sum()))
        
        st.divider()
        
        # 테이블 (세무사 급여대장 형식)
        st.markdown("#### 📋 급여 상세 내역")
        
        # 2026년 01월분 급여대장 형식 (세무사 급여대장)
        detailed_data = []
        for idx, payroll in enumerate(monthly_payroll, 1):
            allowances = payroll.get('allowances', {})
            
            # 기본 수당 추출
            meal_allowance = allowances.get('식대', 0)
            transport_allowance = allowances.get('교통비', 0)
            overtime_allowance = allowances.get('연장근로수당', 0) + allowances.get('야간근로수당', 0) + allowances.get('휴일근로수당', 0)
            other_allowances = sum([v for k, v in allowances.items() if k not in ['식대', '교통비', '연장근로수당', '야간근로수당', '휴일근로수당']])
            
            row = {
                '번호': idx,
                '성명': payroll['name'],
                '부서': payroll['department'],
                '기본급': payroll['base_salary'],
                '식대': meal_allowance,
                '연장근로수당': overtime_allowance,
                '지급합계': payroll['base_salary'] + payroll['total_allowance'],
                '국민연금': payroll.get('national_pension', 0),
                '건강보험': payroll.get('health_insurance', 0),
                '고용보험': payroll.get('employment_insurance', 0),
                '소득세': payroll.get('income_tax', 0),
                '지방소득세': payroll.get('local_tax', 0),
                '공제합계': payroll['total_deduction'],
                '실수령액': payroll['net_pay']
            }
            detailed_data.append(row)
        
        detailed_df = pd.DataFrame(detailed_data)
        
        # 숫자 포맷 적용
        display_df = detailed_df.copy()
        for col in display_df.columns:
            if col not in ['번호', '성명', '부서']:
                display_df[col] = display_df[col].apply(C.format_currency)
        
        st.dataframe(
            display_df,
            use_container_width=True,
            hide_index=True
        )
        
        # 지급 상태 변경
        st.divider()
        st.markdown("### 💳 지급 상태 관리")
        
        unpaid_employees = [p for p in monthly_payroll if p['paid_status'] == '미지급']
        
        if unpaid_employees:
            st.info(f"💡 미지급 직원: {len(unpaid_employees)}명")
            
            for emp in unpaid_employees:
                col1, col2, col3 = st.columns([2, 2, 1])
                
                with col1:
                    st.write(f"**{emp['name']}** ({emp['department']})")
                with col2:
                    st.write(f"실수령액: {C.format_currency(emp['net_pay'])}")
                with col3:
                    if st.button(f"✅ 지급완료", key=f"pay_{emp['emp_id']}", use_container_width=True):
                        if update_paid_status(emp['emp_id'], year_month, '지급완료'):
                            st.success(f"✅ {emp['name']}님의 급여가 지급 완료 처리되었습니다!")
                            st.rerun()
                        else:
                            st.error("❌ 상태 변경 실패")
        else:
            st.success("✅ 모든 급여가 지급 완료되었습니다!")
        
        # 엑셀 다운로드
        st.divider()
        
        # 세무사 급여대장 형식으로 엑셀 생성
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            detailed_df.to_excel(writer, sheet_name='급여대장', index=False)
        
        st.download_button(
            label="📥 급여대장 엑셀 다운로드",
            data=buffer.getvalue(),
            file_name=f"급여대장_{year_month}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            type="primary"
        )

# ============================================================
# 시간외 수당
# ============================================================

elif menu == "⏰ 시간외 수당":
    st.subheader("⏰ 시간외 수당 관리")
    
    # 직원 선택
    employees = get_all_employees(active_only=True)
    employee_options = {f"{emp['name']} ({emp['department']})": emp for emp in employees}
    
    selected = st.selectbox("👤 직원 선택", list(employee_options.keys()))
    
    if selected:
        employee = employee_options[selected]
        emp_id = employee['emp_id']
        
        # 급여 설정 확인
        setting = get_payroll_setting(emp_id)
        
        if not setting:
            st.warning(f"""
            ⚠️ **{employee['name']}님의 급여 설정이 없습니다!**
            
            시간급을 계산할 수 없습니다.
            
            📍 http://localhost:8000 → 💰 급여 정보 관리
            """)
            st.stop()
        else:
            base_salary = setting['base_salary']
            meal_allowance = setting.get('allowances', {}).get('식대', 0)
            fixed_ot_hours = setting.get('fixed_ot_hours', 0)
            is_inclusive_wage = setting.get('is_inclusive_wage', False)
            
            # 통상시급 계산 (기본급 + 식대 포함)
            hourly_wage = calculate_hourly_wage(base_salary, meal_allowance)
            
            st.divider()
            
            # 통상임금 정보 표시
            st.markdown("### 💰 통상임금 정보")
            col_info1, col_info2, col_info3 = st.columns(3)
            with col_info1:
                st.metric("기본급", C.format_currency(base_salary))
            with col_info2:
                st.metric("식대", C.format_currency(meal_allowance), help="통상임금에 포함")
            with col_info3:
                st.metric("통상임금", C.format_currency(base_salary + meal_allowance))
            
            st.info(f"""
            💡 **통상시급**: {C.format_currency(hourly_wage)}원
            - 계산식: (기본급 {C.format_currency(base_salary)} + 식대 {C.format_currency(meal_allowance)}) ÷ {C.COMMON_WAGE_DIVISOR}시간
            - 통상임금 = 기본급 + 식대
            """)
            
            # 고정 OT 정보 표시
            if is_inclusive_wage and fixed_ot_hours > 0:
                st.info(f"🔵 **포괄임금제 적용** (고정 OT: {fixed_ot_hours}시간)")
                st.caption("※ 고정 OT를 초과한 시간외 근무만 추가 수당으로 지급됩니다.")
            
            # 시간외 수당 계산
            with st.form("overtime_form"):
                st.markdown("### 📝 시간외 근무 등록")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    work_date = st.date_input("근무 날짜", value=datetime.now())
                
                with col2:
                    overtime_type = st.selectbox("근무 유형", ["연장", "야간", "휴일"])
                
                st.divider()
                st.markdown("#### 입력 방법 선택")
                input_method = st.radio(
                    "입력 방법",
                    ["⏰ 시간 입력", "💰 금액 입력"],
                    horizontal=True,
                    help="시간을 직접 입력하거나, 금액을 입력하면 시간이 자동 계산됩니다."
                )
                
                hours = 0.0
                overtime_pay = 0.0
                
                if input_method == "⏰ 시간 입력":
                    hours = st.number_input(
                        "실제 근무 시간",
                        min_value=0.0,
                        max_value=24.0,
                        value=2.0,
                        step=0.5,
                        help="이번 달 누적 시간외 근무 시간"
                    )
                    
                    # 고정 OT 초과분 계산
                    if is_inclusive_wage and fixed_ot_hours > 0:
                        monthly_overtime_logs = get_monthly_overtime(emp_id, year_month)
                        total_overtime_this_month = sum([log['hours'] for log in monthly_overtime_logs])
                        remaining_fixed_ot = max(0, fixed_ot_hours - total_overtime_this_month)
                        
                        if hours <= remaining_fixed_ot:
                            billable_hours = 0
                            st.warning(f"⚠️ **고정 OT 범위 내 근무** - 추가 수당 없음")
                        else:
                            billable_hours = hours - remaining_fixed_ot
                    else:
                        billable_hours = hours
                    
                    # 수당 계산
                    overtime_pay = calculate_overtime_pay(base_salary, meal_allowance, billable_hours, overtime_type)
                    
                else:  # 금액 입력
                    overtime_pay = st.number_input(
                        "연장근로수당 금액",
                        min_value=0,
                        value=0,
                        step=1000,
                        help="금액을 입력하면 시간이 자동 계산됩니다."
                    )
                    
                    if overtime_pay > 0:
                        # 시간 역산
                        hours = calculate_ot_hours_from_pay(base_salary, meal_allowance, overtime_pay, overtime_type)
                        
                        # 고정 OT 초과분 계산
                        if is_inclusive_wage and fixed_ot_hours > 0:
                            monthly_overtime_logs = get_monthly_overtime(emp_id, year_month)
                            total_overtime_this_month = sum([log['hours'] for log in monthly_overtime_logs])
                            remaining_fixed_ot = max(0, fixed_ot_hours - total_overtime_this_month)
                            
                            if hours <= remaining_fixed_ot:
                                billable_hours = 0
                                st.warning(f"⚠️ **고정 OT 범위 내 근무** - 추가 수당 없음")
                            else:
                                billable_hours = hours - remaining_fixed_ot
                                overtime_pay = calculate_overtime_pay(base_salary, meal_allowance, billable_hours, overtime_type)
                        else:
                            billable_hours = hours
                        
                        st.success(f"✅ 계산된 연장근로시간: **{hours}시간**")
                
                # 수당 미리보기
                if hours > 0 or overtime_pay > 0:
                    st.markdown(f"""
                    ### 💰 예상 시간외 수당
                    
                    - 통상시급: {C.format_currency(hourly_wage)}
                    - 근무 시간: {hours}시간
                    - 가산율: {C.OVERTIME_RATE if overtime_type == "연장" else C.WORK_TIME.get(f"{overtime_type}근로", {}).get("가산율", 1.5)} 배
                    - **시간외 수당**: {C.format_currency(overtime_pay)}
                    """)
                
                submitted = st.form_submit_button("💾 등록", use_container_width=True, type="primary")
                
                if submitted:
                    # 주 52시간 초과 검증
                    # 해당 주의 총 근로시간 계산 (기본 40시간 + 이번 달 누적 시간외)
                    monthly_overtime_logs = get_monthly_overtime(emp_id, year_month)
                    total_overtime_this_month = sum([log['hours'] for log in monthly_overtime_logs]) + hours
                    
                    # 주당 평균 시간외 근무 (월 4주 기준)
                    weekly_avg_overtime = total_overtime_this_month / 4
                    total_weekly_hours = 40 + weekly_avg_overtime
                    
                    if total_weekly_hours > 52:
                        st.warning(f"""
                        ⚠️ **주 52시간 초과 경고**
                        
                        - 이번 달 누적 시간외: {total_overtime_this_month:.1f}시간
                        - 주당 평균 근로시간: {total_weekly_hours:.1f}시간
                        - 초과 시간: {total_weekly_hours - 52:.1f}시간
                        
                        💡 근로기준법 제53조에 따라 주 최대 근로시간은 52시간입니다.
                        (연장근로 포함: 기본 40시간 + 연장 12시간)
                        """)
                        
                        # 그래도 등록은 가능하도록 (경고만)
                        if st.button("⚠️ 확인했습니다. 등록하기", type="secondary"):
                            if add_overtime_log(emp_id, str(work_date), overtime_type, hours, overtime_pay):
                                st.success("✅ 시간외 근무가 등록되었습니다!")
                                st.rerun()
                            else:
                                st.error("❌ 등록 실패")
                    else:
                        if add_overtime_log(emp_id, str(work_date), overtime_type, hours, overtime_pay):
                            show_success("시간외 근무가 등록되었습니다!")
                            st.rerun()
                        else:
                            st.error("❌ 등록 실패")
            
            # 이번 달 시간외 근무 내역
            st.divider()
            st.markdown(f"### 📋 {year_month} 시간외 근무 내역")
            
            overtime_logs = get_monthly_overtime(emp_id, year_month)
            
            if overtime_logs:
                df = pd.DataFrame(overtime_logs)
                df['overtime_pay'] = df['overtime_pay'].apply(C.format_currency)
                df = df.rename(columns={
                    'work_date': '근무일',
                    'overtime_type': '유형',
                    'hours': '시간',
                    'overtime_pay': '수당'
                })
                
                st.dataframe(df, use_container_width=True, hide_index=True)
                
                total_overtime_pay = sum([log['overtime_pay'] for log in overtime_logs])
                st.metric("**총 시간외 수당**", C.format_currency(total_overtime_pay))
            else:
                st.info("💡 이번 달 시간외 근무 내역이 없습니다.")

# ============================================================
# 연차 관리
# ============================================================

elif menu == "📅 연차 관리":
    st.subheader("📅 연차 관리")
    
    st.info("""
    💡 **연차 계산 기준**
    - 1년 미만: 월 1개씩 발생
    - 1년 이상: 년 15개
    - 3년 이상: 매 2년마다 1개 추가 (최대 25개)
    """)
    
    # 직원 선택
    employees = get_all_employees(active_only=True)
    employee_options = {f"{emp['name']} ({emp['department']})": emp for emp in employees}
    
    selected = st.selectbox("👤 직원 선택", list(employee_options.keys()))
    
    if selected:
        employee = employee_options[selected]
        emp_id = employee['emp_id']
        
        # 입사일 확인
        if not employee.get('hire_date'):
            st.warning("⚠️ 입사일 정보가 없습니다.")
        else:
            hire_date = datetime.strptime(employee['hire_date'], "%Y-%m-%d")
            
            # 연차 발생 일수 계산
            current_year = datetime.now().year
            annual_leave_days = AnnualLeaveCalculator.calculate_annual_leave_days(hire_date)
            
            # DB에 연차 정보가 없으면 초기화
            init_annual_leave_if_not_exists(emp_id, current_year, annual_leave_days)
            
            # 연차 정보 조회
            leave_info = get_annual_leave(emp_id, current_year)
            
            if not leave_info:
                st.error("❌ 연차 정보를 불러올 수 없습니다.")
            else:
                st.divider()
                st.markdown(f"### 📊 {employee['name']}님의 {current_year}년 연차 정보")
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("입사일", hire_date.strftime("%Y-%m-%d"))
                with col2:
                    work_years = (datetime.now() - hire_date).days / 365.25
                    st.metric("근속 연수", f"{work_years:.1f}년")
                with col3:
                    st.metric("📅 발생 연차", f"{leave_info['total_days']}일")
                with col4:
                    st.metric("✅ 사용 연차", f"{leave_info['used_days']}일")
                
                # 남은 연차 강조 표시
                remaining_days = leave_info['remaining_days']
                if remaining_days < 5:
                    color = "#ff4b4b"
                elif remaining_days < 10:
                    color = "#ffa500"
                else:
                    color = "#00cc00"
                
                st.markdown(f"""
                <div style="background-color: {color}; padding: 1rem; border-radius: 8px; text-align: center; margin: 1rem 0;">
                    <h2 style="color: white; margin: 0;">💚 남은 연차: {remaining_days}일</h2>
                </div>
                """, unsafe_allow_html=True)
                
                # 연차 사용 등록
                st.divider()
                st.markdown("### 📝 연차 사용 등록")
                
                with st.form("annual_leave_usage_form"):
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        leave_date = st.date_input(
                            "연차 사용일",
                            value=datetime.now(),
                            help="연차를 사용한 날짜"
                        )
                    
                    with col2:
                        days_options = [0.5, 1.0, 2.0, 3.0, 4.0, 5.0]
                        days = st.selectbox(
                            "사용 일수",
                            options=days_options,
                            index=1,
                            help="0.5일 = 반차"
                        )
                    
                    with col3:
                        leave_type = st.selectbox(
                            "휴가 유형",
                            ["연차", "반차", "병가", "경조사", "공가", "기타"]
                        )
                    
                    reason = st.text_input(
                        "사유 (선택)",
                        help="연차 사용 사유를 입력하세요"
                    )
                    
                    submitted = st.form_submit_button("💾 연차 사용 등록", use_container_width=True, type="primary")
                    
                    if submitted:
                        if days > remaining_days:
                            st.error(f"❌ 남은 연차({remaining_days}일)보다 많이 사용할 수 없습니다!")
                        else:
                            if add_annual_leave_usage(emp_id, str(leave_date), days, leave_type, reason):
                                show_success(f"{employee['name']}님의 연차 사용이 등록되었습니다!")
                                st.rerun()
                            else:
                                st.error("❌ 연차 사용 등록 실패")
                
                # 연차 사용 이력
                st.divider()
                st.markdown(f"### 📋 {current_year}년 연차 사용 이력")
                
                usage_history = get_annual_leave_usage(emp_id, current_year)
                
                if usage_history:
                    df = pd.DataFrame(usage_history)
                    df = df.rename(columns={
                        'leave_date': '사용일',
                        'days': '일수',
                        'leave_type': '유형',
                        'reason': '사유'
                    })
                    
                    # 사유가 None인 경우 빈 문자열로 변환
                    df['사유'] = df['사유'].fillna('')
                    
                    st.dataframe(
                        df[['사용일', '일수', '유형', '사유']],
                        use_container_width=True,
                        hide_index=True
                    )
                    
                    total_used = sum([h['days'] for h in usage_history])
                    st.info(f"💡 총 사용 연차: **{total_used}일** / 발생 연차: **{leave_info['total_days']}일**")
                else:
                    st.info("💡 아직 연차 사용 이력이 없습니다.")

# ============================================================
# 급여명세서 출력
# ============================================================

elif menu == "📄 급여명세서 출력":
    st.subheader(f"📄 {year_month} 급여명세서")
    
    # 직원 선택
    employees = get_all_employees(active_only=True)
    employee_options = {f"{emp['name']} ({emp['department']})": emp for emp in employees}
    
    selected = st.selectbox("👤 직원 선택", list(employee_options.keys()))
    
    if selected:
        employee = employee_options[selected]
        emp_id = employee['emp_id']
        
        # 해당 월 급여 이력 조회 (수정: 직접 해당 월 조회)
        payroll = get_payroll_history(emp_id, year_month)
        
        if not payroll:
            st.warning(f"⚠️ {year_month} 급여 이력이 없습니다. '💰 월별 급여 계산' 메뉴에서 먼저 계산하세요.")
        else:
            
            st.divider()
            
            # 회사 정보
            company = get_company_profile()
            company_name = company['company_name'] if company else "회사명"
            
            # 급여 설정 불러오기
            setting = get_payroll_setting(emp_id)
            
            # 통상시급 계산 (기본급 + 식대 포함)
            base_salary = setting.get('base_salary', 0) if setting else 0
            meal_allowance = setting.get('allowances', {}).get('식대', 0) if setting else 0
            calculated_hourly_wage = calculate_hourly_wage(base_salary, meal_allowance) if setting else 0
            
            # 수당 내역 HTML 생성
            allowances_html = ""
            if payroll.get('allowances'):
                for name, amount in payroll.get('allowances', {}).items():
                    if amount > 0:
                        tax_status = "비과세" if amount <= C.TAX_FREE_LIMITS.get(name, 0) else "과세"
                        allowances_html += f"""
                        <tr>
                            <td style="padding: 0.5rem; border: 1px solid #000;">{name}</td>
                            <td style="padding: 0.5rem; border: 1px solid #000;">{tax_status}</td>
                            <td style="padding: 0.5rem; border: 1px solid #000; text-align: right;">{amount:,}원</td>
                        </tr>
                        """
            
            # 포괄임금제 안내
            inclusive_wage_info = ""
            if setting and setting.get('is_inclusive_wage'):
                inclusive_wage_info = f'<p style="font-size: 0.9em; margin: 0.3rem 0; color: #ff6600;"><strong>※ 포괄임금제 적용:</strong> 고정 OT {setting.get("fixed_ot_hours", 0)}시간 포함</p>'
            
            # 급여명세서 HTML (표준양식 참고)
            payslip_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @media print {{
                        body {{ margin: 0; padding: 0; }}
                        .container {{ border: 2px solid #000; page-break-inside: avoid; }}
                        @page {{ size: A4; margin: 1cm; }}
                    }}
                    body {{
                        font-family: 'Malgun Gothic', 'Apple SD Gothic Neo', sans-serif;
                        padding: 10px;
                        background: white;
                        margin: 0;
                    }}
                    .container {{
                        max-width: 800px;
                        margin: 0 auto;
                        border: 2px solid #000;
                        background: white;
                    }}
                    h1 {{
                        text-align: center;
                        margin: 0;
                        padding: 1rem 0;
                        font-size: 1.8em;
                        font-weight: bold;
                        border-bottom: 2px solid #000;
                    }}
                    .info-section {{
                        text-align: right;
                        padding: 0.5rem 1rem;
                        font-size: 0.9em;
                        border-bottom: 1px solid #000;
                    }}
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                    }}
                    td {{
                        border: 1px solid #000;
                        padding: 0.4rem;
                        font-size: 0.9em;
                    }}
                    .header-cell {{
                        background-color: #e8e8e8;
                        font-weight: bold;
                        text-align: center;
                    }}
                    .section-header {{
                        background-color: #d0d0d0;
                        font-weight: bold;
                        text-align: center;
                        padding: 0.5rem;
                    }}
                    .amount {{
                        text-align: right;
                        font-weight: bold;
                    }}
                    .total-row {{
                        background-color: #f5f5f5;
                        font-weight: bold;
                    }}
                    .net-pay-row {{
                        background-color: #fff5cc;
                        font-weight: bold;
                        font-size: 1.1em;
                    }}
                    .calc-section {{
                        margin-top: 1rem;
                    }}
                    .calc-header {{
                        background-color: #e0e0e0;
                        font-weight: bold;
                        text-align: center;
                        padding: 0.5rem;
                    }}
                    .notice {{
                        padding: 0.5rem 1rem;
                        font-size: 0.8em;
                        color: #666;
                        border-top: 1px solid #000;
                        margin-top: 1rem;
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>임 금 명 세 서</h1>
                    
                    <div class="info-section">
                        지급일: {year_month}-{C.DEFAULT_PAYDAY}
                    </div>
                    
                    <!-- 기본 정보 (인적 사항) -->
                    <table>
                        <tr>
                            <td class="header-cell" style="width: 15%;">성명</td>
                            <td style="width: 35%;">{employee['name']}</td>
                            <td class="header-cell" style="width: 15%;">사번</td>
                            <td style="width: 35%;">{employee.get('emp_id', '-')}</td>
                        </tr>
                        <tr>
                            <td class="header-cell">생년월일</td>
                            <td>{employee.get('resident_number', '')[:6] if employee.get('resident_number') else '-'}</td>
                            <td class="header-cell">귀속년월</td>
                            <td>{year_month}</td>
                        </tr>
                        <tr>
                            <td class="header-cell">부서</td>
                            <td>{employee['department']}</td>
                            <td class="header-cell">직급</td>
                            <td>{employee['position']}</td>
                        </tr>
                    </table>
                    
                    <!-- 세부 내역 (지급/공제 좌우 배치) -->
                    <table>
                        <tr>
                            <td colspan="4" class="section-header">세부 내역</td>
                        </tr>
                        <tr>
                            <td colspan="2" class="header-cell">지 급</td>
                            <td colspan="2" class="header-cell">공 제</td>
                        </tr>
                        <tr>
                            <td class="header-cell" style="width: 20%;">임금 항목</td>
                            <td class="header-cell" style="width: 30%;">지급 금액(원)</td>
                            <td class="header-cell" style="width: 20%;">공제 항목</td>
                            <td class="header-cell" style="width: 30%;">공제 금액(원)</td>
                        </tr>
                        <tr>
                            <td>기본급</td>
                            <td class="amount">{payroll['base_salary']:,}</td>
                            <td>소득세</td>
                            <td class="amount">{payroll['income_tax']:,}</td>
                        </tr>"""
            
            # 수당 및 공제 항목 동적 생성
            allowance_items = []
            
            # 수당 상세 추가 (식대, 연장근로수당 등)
            if payroll.get('allowances'):
                for name, amount in payroll.get('allowances', {}).items():
                    if amount > 0:
                        allowance_items.append((name, amount))
            
            # 공제 항목 (소득세는 이미 첫 줄에 있으므로 나머지만)
            deduction_items = []
            
            # 0원이 아닌 공제 항목만 추가
            if payroll.get('national_pension', 0) > 0:
                deduction_items.append(("국민연금", payroll['national_pension']))
            if payroll.get('health_insurance', 0) > 0:
                deduction_items.append(("건강보험", payroll['health_insurance']))
            if payroll.get('employment_insurance', 0) > 0:
                deduction_items.append(("고용보험", payroll['employment_insurance']))
            if payroll.get('longterm_care', 0) > 0:
                deduction_items.append(("장기요양보험", payroll['longterm_care']))
            if payroll.get('local_tax', 0) > 0:
                deduction_items.append(("지방소득세", payroll['local_tax']))
            
            # 최대 줄 수 계산
            max_rows = max(len(allowance_items), len(deduction_items))
            
            for i in range(max_rows):
                payslip_html += "<tr>"
                # 지급 항목 (수당)
                if i < len(allowance_items):
                    payslip_html += f"<td>{allowance_items[i][0]}</td><td class='amount'>{allowance_items[i][1]:,}</td>"
                else:
                    payslip_html += "<td></td><td></td>"
                # 공제 항목
                if i < len(deduction_items):
                    item = deduction_items[i]
                    payslip_html += f"<td>{item[0]}</td><td class='amount'>{item[1]:,}</td>"
                else:
                    payslip_html += "<td></td><td></td>"
                payslip_html += "</tr>"
            
            # 총 지급액 계산 (기본급 + 모든 수당)
            total_payment = payroll['base_salary'] + sum(payroll.get('allowances', {}).values())
            
            payslip_html += f"""
                        <tr class="total-row">
                            <td>지급액 계</td>
                            <td class="amount">{total_payment:,}</td>
                            <td>공제액 계</td>
                            <td class="amount">{payroll['total_deduction']:,}</td>
                        </tr>
                        <tr class="net-pay-row">
                            <td colspan="3" style="text-align: center;">실수령액(원)</td>
                            <td class="amount" style="font-size: 1.2em;">{payroll['net_pay']:,}</td>
                        </tr>
                    </table>
                    
                    <!-- 근로시간 및 계산방법 -->
                    <table class="calc-section">
                        <tr>
                            <td class="header-cell" style="width: 20%;">연장근로시간수</td>
                            <td style="width: 13%;" class="amount">{setting.get('fixed_ot_hours', 0) if setting and setting.get('is_inclusive_wage') else '-'}</td>
                            <td class="header-cell" style="width: 20%;">야간근로시간수</td>
                            <td style="width: 13%;" class="amount">-</td>
                            <td class="header-cell" style="width: 20%;">휴일근로시간수</td>
                            <td style="width: 14%;" class="amount">-</td>
                        </tr>
                        <tr>
                            <td class="header-cell">통상시급(원)</td>
                            <td class="amount">{calculated_hourly_wage:,.0f}</td>
                            <td colspan="4"></td>
                        </tr>
                    </table>
                    
                    <!-- 계산 방법 (실제 적용된 값 표시) -->
                    <table class="calc-section">
                        <tr>
                            <td colspan="2" class="calc-header">계산 방법 (해당 직원 적용 내역)</td>
                        </tr>
                        <tr>
                            <td class="header-cell" style="width: 30%;">구분</td>
                            <td class="header-cell" style="width: 70%;">산출식 (실제 적용 값)</td>
                        </tr>"""
            
            # 실제 적용된 보험 및 세금 계산 표시
            if payroll.get('national_pension', 0) > 0:
                payslip_html += f"""
                        <tr>
                            <td class="header-cell">국민연금</td>
                            <td>{payroll['taxable_amount']:,}원 (과세급여) × 4.75% = {payroll['national_pension']:,}원</td>
                        </tr>"""
            
            if payroll.get('health_insurance', 0) > 0:
                payslip_html += f"""
                        <tr>
                            <td class="header-cell">건강보험</td>
                            <td>{payroll['taxable_amount']:,}원 (과세급여) × 3.60% = {payroll['health_insurance']:,}원</td>
                        </tr>"""
            
            if payroll.get('longterm_care', 0) > 0:
                payslip_html += f"""
                        <tr>
                            <td class="header-cell">장기요양보험</td>
                            <td>{payroll['health_insurance']:,}원 (건강보험료) × 13.14% = {payroll['longterm_care']:,}원</td>
                        </tr>"""
            
            if payroll.get('employment_insurance', 0) > 0:
                payslip_html += f"""
                        <tr>
                            <td class="header-cell">고용보험</td>
                            <td>{payroll['taxable_amount']:,}원 (과세급여) × 0.9% = {payroll['employment_insurance']:,}원</td>
                        </tr>"""
            
            payslip_html += f"""
                        <tr>
                            <td class="header-cell">소득세</td>
                            <td>간이세액표 기준 (본인 1명) = {payroll['income_tax']:,}원</td>
                        </tr>
                        <tr>
                            <td class="header-cell">지방소득세</td>
                            <td>{payroll['income_tax']:,}원 (소득세) × 10% = {payroll['local_tax']:,}원</td>
                        </tr>"""
            
            # 연장근로수당이 있는 경우
            if setting and setting.get('is_inclusive_wage') and setting.get('fixed_ot_hours', 0) > 0:
                fixed_ot_hours = setting.get('fixed_ot_hours', 0)
                fixed_ot_amount = setting.get('fixed_ot_amount', 0)
                payslip_html += f"""
                        <tr>
                            <td class="header-cell">연장근로수당</td>
                            <td>{calculated_hourly_wage:,.0f}원 (통상시급) × {fixed_ot_hours:.1f}시간 × {C.OVERTIME_RATE}배 = {fixed_ot_amount:,}원</td>
                        </tr>"""
            
            payslip_html += """
                    </table>
                    
                    <div class="notice">
                        ※ 근로기준법 제48조에 따라 임금명세서를 교부합니다. | 2026년 최저시급: 10,320원 | 법정근로시간: 주 40시간<br>
                        ※ 발행: {datetime.now().strftime("%Y년 %m월 %d일")} | {company_name}
                        {' | 포괄임금제 적용 (고정 OT ' + str(setting.get('fixed_ot_hours', 0)) + '시간)' if setting and setting.get('is_inclusive_wage') else ''}<br>
                        ※ <strong>해당 명세서는 2026년 개정 근로기준법 및 사회보험 요율을 준수합니다.</strong>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # HTML 렌더링
            st.markdown("### 📄 급여명세서 미리보기")
            components.html(payslip_html, height=1200, scrolling=True)
            
            st.divider()
            
            # 다운로드 옵션
            st.markdown("### 📥 다운로드")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # 워드(DOCX) 다운로드
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    import io
                    
                    # DOCX 문서 생성
                    doc = Document()
                    
                    # 제목
                    title = doc.add_heading('급 여 명 세 서', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 기본 정보
                    doc.add_paragraph(f"귀속년월: {year_month}  |  지급일: {year_month}-{C.DEFAULT_PAYDAY}")
                    doc.add_paragraph(f"성명: {employee['name']}  |  사번: {employee.get('emp_id', '-')}  |  생년월일: {employee.get('resident_number', '')[:6] if employee.get('resident_number') else '-'}")
                    doc.add_paragraph(f"부서: {employee['department']}  |  직급: {employee['position']}")
                    doc.add_paragraph("")
                    
                    # 지급/공제 내역 표
                    doc.add_heading('지급 및 공제 내역', level=2)
                    
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Light Grid Accent 1'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '지급 항목'
                    hdr_cells[1].text = '지급 금액'
                    hdr_cells[2].text = '공제 항목'
                    hdr_cells[3].text = '공제 금액'
                    
                    # 지급 및 공제 항목 (세무사 급여대장 형식)
                    pay_items = [('기본급', payroll['base_salary'])]
                    
                    # 개별 수당 추가
                    if payroll.get('allowances'):
                        for name, amount in payroll.get('allowances', {}).items():
                            if amount > 0:
                                pay_items.append((name, amount))
                    
                    deduction_items = [
                        ('국민연금', payroll.get('national_pension', 0)),
                        ('건강보험', payroll.get('health_insurance', 0)),
                        ('장기요양', payroll.get('longterm_care', 0)),
                        ('고용보험', payroll.get('employment_insurance', 0)),
                        ('소득세', payroll.get('income_tax', 0)),
                        ('지방소득세', payroll.get('local_tax', 0))
                    ]
                    
                    # 행별로 지급/공제 동시 표시
                    max_rows = max(len(pay_items), len(deduction_items))
                    for i in range(max_rows):
                        row_cells = table.add_row().cells
                        
                        if i < len(pay_items):
                            row_cells[0].text = pay_items[i][0]
                            row_cells[1].text = f"{pay_items[i][1]:,.0f}원"
                        else:
                            row_cells[0].text = ''
                            row_cells[1].text = ''
                        
                        if i < len(deduction_items):
                            row_cells[2].text = deduction_items[i][0]
                            row_cells[3].text = f"{deduction_items[i][1]:,.0f}원"
                        else:
                            row_cells[2].text = ''
                            row_cells[3].text = ''
                    
                    # 합계 행
                    row_cells = table.add_row().cells
                    row_cells[0].text = '총 지급액'
                    row_cells[1].text = f"{payroll['base_salary'] + payroll['total_allowance']:,.0f}원"
                    row_cells[2].text = '총 공제액'
                    row_cells[3].text = f"{payroll['total_deduction']:,.0f}원"
                    
                    # 실수령액
                    doc.add_paragraph("")
                    p = doc.add_paragraph()
                    p.add_run('실수령액: ').bold = True
                    p.add_run(f"{payroll['net_pay']:,.0f}원").bold = True
                    p.runs[1].font.size = Pt(14)
                    
                    # 고정 OT 정보 추가
                    if setting and setting.get('is_inclusive_wage'):
                        doc.add_paragraph("")
                        doc.add_paragraph(f"※ 포괄임금제 적용 (고정 OT {setting.get('fixed_ot_hours', 0)}시간)")
                    
                    # 저장
                    docx_buffer = io.BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    
                    st.download_button(
                        label="📘 워드 다운로드",
                        data=docx_buffer.getvalue(),
                        file_name=f"급여명세서_{employee['name']}_{year_month}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        help="워드 파일로 다운로드하여 편집 가능",
                        use_container_width=True
                    )
                except ImportError:
                    st.button(
                        "📘 워드 다운로드",
                        disabled=True,
                        help="python-docx 라이브러리 설치 필요",
                        use_container_width=True
                    )
                except Exception as e:
                    st.button(
                        "📘 워드 다운로드",
                        disabled=True,
                        help=f"워드 생성 오류: {str(e)}",
                        use_container_width=True
                    )
            
            with col2:
                # 엑셀 다운로드
                try:
                    from openpyxl import Workbook
                    from openpyxl.styles import Font, Alignment, PatternFill
                    import io
                    
                    # 워크북 생성
                    wb = Workbook()
                    ws = wb.active
                    ws.title = "급여명세서"
                    
                    # 제목
                    ws.merge_cells('A1:B1')
                    ws['A1'] = '급 여 명 세 서'
                    ws['A1'].font = Font(size=16, bold=True)
                    ws['A1'].alignment = Alignment(horizontal='center')
                    
                    # 기본 정보
                    ws['A3'] = '귀속년월'
                    ws['B3'] = year_month
                    ws['A4'] = '성명'
                    ws['B4'] = employee['name']
                    ws['A5'] = '사번'
                    ws['B5'] = employee.get('emp_id', '-')
                    ws['A6'] = '생년월일'
                    ws['B6'] = employee.get('resident_number', '')[:6] if employee.get('resident_number') else '-'
                    ws['A7'] = '부서'
                    ws['B7'] = employee['department']
                    ws['A8'] = '직급'
                    ws['B8'] = employee['position']
                    ws['A9'] = '지급일'
                    ws['B9'] = f"{year_month}-{C.DEFAULT_PAYDAY}"
                    
                    # 지급/공제 내역 (세무사 급여대장 형식)
                    ws['A11'] = '지급 항목'
                    ws['B11'] = '지급 금액'
                    ws['C11'] = '공제 항목'
                    ws['D11'] = '공제 금액'
                    
                    # 헤더 스타일
                    for col in ['A11', 'B11', 'C11', 'D11']:
                        ws[col].font = Font(bold=True)
                        ws[col].fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
                        ws[col].alignment = Alignment(horizontal='center')
                    
                    # 지급 항목 준비
                    pay_items = [('기본급', payroll['base_salary'])]
                    if payroll.get('allowances'):
                        for name, amount in payroll.get('allowances', {}).items():
                            if amount > 0:
                                pay_items.append((name, amount))
                    
                    # 공제 항목
                    deduction_items = [
                        ('국민연금', payroll.get('national_pension', 0)),
                        ('건강보험', payroll.get('health_insurance', 0)),
                        ('장기요양', payroll.get('longterm_care', 0)),
                        ('고용보험', payroll.get('employment_insurance', 0)),
                        ('소득세', payroll.get('income_tax', 0)),
                        ('지방소득세', payroll.get('local_tax', 0))
                    ]
                    
                    # 데이터 입력
                    row = 12
                    max_rows = max(len(pay_items), len(deduction_items))
                    
                    for i in range(max_rows):
                        # 지급 항목
                        if i < len(pay_items):
                            ws[f'A{row}'] = pay_items[i][0]
                            ws[f'B{row}'] = pay_items[i][1]
                            ws[f'B{row}'].number_format = '#,##0'
                        
                        # 공제 항목
                        if i < len(deduction_items):
                            ws[f'C{row}'] = deduction_items[i][0]
                            ws[f'D{row}'] = deduction_items[i][1]
                            ws[f'D{row}'].number_format = '#,##0'
                        
                        row += 1
                    
                    # 합계 행
                    ws[f'A{row}'] = '총 지급액'
                    ws[f'B{row}'] = payroll['base_salary'] + payroll['total_allowance']
                    ws[f'B{row}'].number_format = '#,##0'
                    ws[f'B{row}'].font = Font(bold=True)
                    
                    ws[f'C{row}'] = '총 공제액'
                    ws[f'D{row}'] = payroll['total_deduction']
                    ws[f'D{row}'].number_format = '#,##0'
                    ws[f'D{row}'].font = Font(bold=True)
                    row += 1
                    
                    # 실수령액
                    ws[f'A{row+1}'] = '실수령액'
                    ws[f'B{row+1}'] = payroll['net_pay']
                    ws[f'B{row+1}'].number_format = '#,##0'
                    ws[f'B{row+1}'].font = Font(bold=True, size=14)
                    ws[f'B{row+1}'].fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                    
                    # 고정 OT 정보
                    if setting and setting.get('is_inclusive_wage'):
                        ws[f'A{row+1}'] = f"※ 포괄임금제 적용 (고정 OT {setting.get('fixed_ot_hours', 0)}시간)"
                    
                    # 열 너비 조정
                    ws.column_dimensions['A'].width = 20
                    ws.column_dimensions['B'].width = 20
                    ws.column_dimensions['C'].width = 20
                    ws.column_dimensions['D'].width = 20
                    
                    # 저장
                    excel_buffer = io.BytesIO()
                    wb.save(excel_buffer)
                    excel_buffer.seek(0)
                    
                    st.download_button(
                        label="📗 엑셀 다운로드",
                        data=excel_buffer.getvalue(),
                        file_name=f"급여명세서_{employee['name']}_{year_month}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        help="엑셀 파일로 다운로드하여 편집 가능",
                        use_container_width=True
                    )
                except Exception as e:
                    st.button(
                        "📗 엑셀 다운로드",
                        disabled=True,
                        help=f"엑셀 생성 오류: {str(e)}",
                        use_container_width=True
                    )
            
            with col3:
                # HTML 파일 다운로드
                st.download_button(
                    label="📄 HTML 다운로드",
                    data=payslip_html.encode('utf-8'),
                    file_name=f"급여명세서_{employee['name']}_{year_month}.html",
                    mime="text/html",
                    help="브라우저에서 열어서 인쇄(Ctrl+P) 가능",
                    use_container_width=True
                )
            
            st.caption("💡 **추천**: HTML 다운로드 후 브라우저에서 인쇄 (서식 완벽 유지) | 워드/엑셀 (편집 가능)")
            
            st.divider()

# ============================================================
# 사이드바 정보
# ============================================================

st.sidebar.divider()
st.sidebar.markdown(f"""
### 💡 시스템 정보

- **버전**: {C.APP_CONFIG['버전']}
- **최저임금 (2026)**: {C.format_currency(C.MINIMUM_WAGE['시급'])}/시간
- **법정근로시간**: 주 {C.WORK_TIME['법정근로시간']['주']}시간
- **국민연금**: {C.INSURANCE_RATES['국민연금']['요율']*100}%
- **건강보험**: {C.INSURANCE_RATES['건강보험']['요율']*100}%
- **고용보험**: {C.INSURANCE_RATES['고용보험']['요율']*100}%
""")

# ============================================================
# 초기화 (앱 시작 시)
# ============================================================

# 데이터베이스 초기화
try:
    init_payroll_tables()
except Exception as e:
    st.error(f"데이터베이스 초기화 실패: {e}")
