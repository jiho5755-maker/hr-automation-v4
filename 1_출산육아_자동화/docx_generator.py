"""
워드(DOCX) 기반 임신 관련 서식 생성기
원본 서식과 100% 동일하게 구현 - 한 페이지 완성
"""

from io import BytesIO
from datetime import date, datetime
from typing import Dict, Optional
import calendar

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import constants as C


def shade_cell(cell, color="D9D9D9"):
    """셀에 배경색 추가"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


# ============================================================
# 2. 임신사유 근로시간 단축 확인서 (한 페이지 완성)
# ============================================================

def create_confirmation_form_docx(employee_info: Dict, employer_info: Dict,
                                   pregnancy_data: Dict, childbirth_data: Dict) -> BytesIO:
    """임신사유 근로시간 단축 확인서 DOCX 생성 - 한 페이지 완성"""
    
    doc = Document()
    
    # 페이지 여백 최소화
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
    
    # 서식 번호
    header_p = doc.add_paragraph()
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(3)
    header_run = header_p.add_run('■ 고용창출장려금·고용안정장려금의 신청 및 지급에 관한 규정 [별지 제22호의2 서식]')
    header_run.font.size = Pt(7)
    header_run.font.name = 'Malgun Gothic'
    
    # 제목
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('임신사유 근로시간 단축에 대한 근로자 확인서')
    title_run.font.size = Pt(13)
    title_run.font.bold = True
    title_run.font.name = 'Malgun Gothic'
    
    # 동의 문구
    consent_p = doc.add_paragraph()
    consent_p.paragraph_format.space_before = Pt(0)
    consent_p.paragraph_format.space_after = Pt(6)
    consent_run = consent_p.add_run(
        '※ 본인은 임신 사유로 근로시간을 단축한 사실을 아래와 같이 확인하며, 위라밸일자리 장려금\n'
        '    (소정근로시간 단축제) 지원을 위한 자료로 활용하는 것에 대해 동의합니다.'
    )
    consent_run.font.size = Pt(6.5)
    consent_run.font.name = 'Malgun Gothic'
    
    # 데이터 준비
    start = pregnancy_data["시작일"]
    end = pregnancy_data["종료일"]
    
    # 주민등록번호 안전 처리
    birth_str = employee_info.get("주민등록번호", "")
    if birth_str and len(birth_str) >= 6:
        birth_str_safe = birth_str[:6]
        if birth_str_safe[0] in ['0', '1', '2']:
            birth_year = f"20{birth_str_safe[:2]}"
        else:
            birth_year = f"19{birth_str_safe[:2]}"
        birth_formatted = f"({birth_year}.{birth_str_safe[2:4]}.{birth_str_safe[4:6]})"
    else:
        birth_formatted = "(0000.00.00)"
    
    phone = employee_info.get("연락처", "010-xxxx-xxxx")
    
    pregnancy_confirm = childbirth_data.get("임신확인일")
    if pregnancy_confirm and isinstance(pregnancy_confirm, date):
        pregnancy_date_str = f"{pregnancy_confirm.strftime('%y.%m.%d')}"
    else:
        pregnancy_date_str = "00.00.00"
    
    # ========================================
    # 기본 정보 테이블 (2행 6열) - 임신일과 충 단축 기간 구분
    # ========================================
    table1 = doc.add_table(rows=2, cols=6)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 첫 번째 행
    row0_cells = table1.rows[0].cells
    shade_cell(row0_cells[0], "E7E6E6")
    shade_cell(row0_cells[2], "E7E6E6")
    shade_cell(row0_cells[4], "E7E6E6")
    
    row0_cells[0].text = '성  명\n(생년월일)'
    row0_cells[1].text = f"{employee_info['이름']} {birth_formatted}"
    row0_cells[2].text = '연락처'
    row0_cells[3].text = phone
    row0_cells[4].text = '임신일'
    row0_cells[5].text = pregnancy_date_str
    
    # 두 번째 행 (앞 4개 셀 병합, 뒤 2개는 "충 단축 기간")
    table1.cell(1, 0).merge(table1.cell(1, 3))
    row1_cells = table1.rows[1].cells
    
    # 병합된 첫 번째 셀은 비워둠 (또는 확인합니다 문구)
    row1_cells[0].text = '○ 근로시간 단축 기간 내 아래와 같이 근로시간을 단축하였음을 확인합니다.'
    
    shade_cell(row1_cells[4], "E7E6E6")
    row1_cells[4].text = '충 단축 기간'
    row1_cells[5].text = f"{start.strftime('%y.%m.%d')}~\n{end.strftime('%y.%m.%d')}"
    
    # 셀 서식
    for row in table1.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                if '○ 근로시간' in para.text:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.0
                for run in para.runs:
                    if '○ 근로시간' in run.text:
                        run.font.size = Pt(6.5)
                    else:
                        run.font.size = Pt(7)
                    run.font.name = 'Malgun Gothic'
    
    # 헤더 볼드 처리
    for i in [0, 2, 4]:
        for run in table1.rows[0].cells[i].paragraphs[0].runs:
            run.font.bold = True
    for run in table1.rows[1].cells[4].paragraphs[0].runs:
        run.font.bold = True
    
    # ========================================
    # 메인 표: 단축 근로 시간 및 이행 (2열, 컴팩트)
    # ========================================
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 왼쪽 셀: "단축 근로 시간 및 이행"
    left_cell = table2.rows[0].cells[0]
    left_cell.width = Cm(1.2)
    left_cell.text = '단축\n근로\n시간\n및\n이행'
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_cell.paragraphs[0].paragraph_format.line_spacing = 0.9
    shade_cell(left_cell, "E7E6E6")
    for run in left_cell.paragraphs[0].runs:
        run.font.size = Pt(7.5)
        run.font.name = 'Malgun Gothic'
        run.font.bold = True
    
    # 오른쪽 셀: 모든 내용 (폰트 크기 축소)
    right_cell = table2.rows[0].cells[1]
    right_cell.width = Cm(15.6)
    
    # 내용을 한 셀에 모두 작성 (줄바꿈 최소화)
    _, last_day_1 = calendar.monthrange(start.year, start.month)
    
    content_lines = [
        f'1회차) {start.year}년 {start.month:02d}월',
        f'  • 해당월 단축 사용 기간                    {start.year}. {start.month:02d}. {start.day:02d}. ~ {start.year}. {start.month:02d}. {last_day_1}.',
        f'  • 단축 후 주당 근로시간                    주당 30시간',
        f'  • 단축 후 근로시간 준수 여부            준수( O ), 미준수(    )',
        '',
        f'2회차) {end.year}년 {end.month:02d}월',
        f'  • 해당월 단축 사용 기간                    {end.year}. {end.month:02d}. 01. ~ {end.year}. {end.month:02d}. {end.day:02d}.',
        f'  • 단축 후 주당 근로시간                    주당 30시간',
        f'  • 단축 후 근로시간 준수 여부            준수( O ), 미준수(    )',
        '',
        f'3회차) 2020년 00월',
        f'  • 해당월 단축 사용 기간                    2020. 00. 00. ~ 2020. 00. 00.',
        f'  • 단축 후 주당 근로시간                    주당 00시간',
        f'  • 단축 후 근로시간 준수 여부            준수(  ), 미준수(  )',
        '',
        '✧ 작성방법',
        '① 1일부터 말일까지를 월 단위로 매월 작성합니다.',
        '② 단축 후 주당 근로시간: 단축 후 근로계약서 또는 근로시간 단축 신청서 상의 (주당)근로시간을 작성합니다.',
        '   * 주15시간 미만, 주30시간을 초과하는 경우 장려금을 지원하지 않음',
        '③ 단축 후의 근로시간을 초과하여 근무한 날이 있는 월은 미준수에 체크(√)합니다.',
        '   * 사업주의 업무지시 또는 연장근로 승인에 따라 단축 후 근무시간을 초과하여 근무한 경우를 말하며,',
        '     미준수한 월은 해당월의 장려금을 지원하지 않음',
    ]
    
    # 첫 번째 paragraph 사용
    right_para = right_cell.paragraphs[0]
    right_para.paragraph_format.space_before = Pt(1)
    right_para.paragraph_format.space_after = Pt(1)
    right_para.paragraph_format.line_spacing = 0.85
    
    for i, line in enumerate(content_lines):
        if i > 0:
            right_para.add_run('\n')
        run = right_para.add_run(line)
        run.font.size = Pt(6)
        run.font.name = 'Malgun Gothic'
    
    # ========================================
    # 부정수급 처벌 표
    # ========================================
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    warning_left = table3.rows[0].cells[0]
    warning_left.width = Cm(1.2)
    warning_left.text = '부정\n수급\n처벌'
    warning_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    warning_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning_left.paragraphs[0].paragraph_format.line_spacing = 0.9
    shade_cell(warning_left, "E7E6E6")
    for run in warning_left.paragraphs[0].runs:
        run.font.size = Pt(7.5)
        run.font.name = 'Malgun Gothic'
        run.font.bold = True
    
    warning_right = table3.rows[0].cells[1]
    warning_right.width = Cm(15.6)
    
    warning_lines = [
        '○ 만약 거짓이나 그 밖의 부정한 방법으로 위라밸일자리 장려금(소정근로시간단축제)을 지급 받고자',
        '사업주와 공모한 경우에는 해당 근로자도 고용보험법령에 의해 처벌될 수 있음을 확인합니다.',
        '',
        '※ 부정수급 적발시 경우, 「고용보험법」 제116조제1항에 따라 사업주와 공모하여 거짓이나 그 밖의 부',
        '정한 방법으로 다음 고용안정사업의 지원금 또는 급여를 받은 자와 공모한 사업주는 각각 5년 이하',
        '의 징역 또는 5천만원 이하의 벌금 가능',
        '',
        '✧ 해당 위라밸일자리 장려금(소정근로시간단축제) 신청에 따른 사실관계, 부정수급 여부 확인 등을 위해',
        '관할 고용노동관서에서 연락하여 확인할 수 있습니다. (이의 사용 불가)',
    ]
    
    warning_para = warning_right.paragraphs[0]
    warning_para.paragraph_format.space_before = Pt(1)
    warning_para.paragraph_format.space_after = Pt(1)
    warning_para.paragraph_format.line_spacing = 0.9
    
    for i, line in enumerate(warning_lines):
        if i > 0:
            warning_para.add_run('\n')
        run = warning_para.add_run(line)
        run.font.size = Pt(5.5)
        run.font.name = 'Malgun Gothic'
    
    # ========================================
    # 동의 문구 및 서명 (컴팩트)
    # ========================================
    agreement = doc.add_paragraph()
    agreement.paragraph_format.space_before = Pt(4)
    agreement.paragraph_format.space_after = Pt(4)
    agreement_run = agreement.add_run('위 내용에 동의하며 기재 내용이 모두 사실임을 확인합니다.')
    agreement_run.font.size = Pt(7)
    agreement_run.font.name = 'Malgun Gothic'
    
    # 근로자 서명
    today = datetime.now()
    worker_date = doc.add_paragraph()
    worker_date.paragraph_format.space_before = Pt(2)
    worker_date.paragraph_format.space_after = Pt(2)
    worker_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    worker_date_run = worker_date.add_run(f'{today.year}년      {today.month:02d}월      {today.day:02d}일')
    worker_date_run.font.size = Pt(8)
    worker_date_run.font.name = 'Malgun Gothic'
    
    worker_sign = doc.add_paragraph()
    worker_sign.paragraph_format.space_before = Pt(2)
    worker_sign.paragraph_format.space_after = Pt(6)
    worker_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    worker_sign_run = worker_sign.add_run(f"확인자(근로자)  {employee_info['이름']}  (서명 또는 인)")
    worker_sign_run.font.size = Pt(8)
    worker_sign_run.font.name = 'Malgun Gothic'
    
    # 사업주 제출 문구
    submit_text = doc.add_paragraph()
    submit_text.paragraph_format.space_before = Pt(2)
    submit_text.paragraph_format.space_after = Pt(4)
    submit_text_run = submit_text.add_run(
        '「고용창출장려금·고용안정장려금의 신청 및 지급에 관한 규정」에 따라 위와 같이 근로자로부터 확인받았으며, 동 확인서를 제출합니다.'
    )
    submit_text_run.font.size = Pt(7)
    submit_text_run.font.name = 'Malgun Gothic'
    
    # 사업주 서명
    employer_date = doc.add_paragraph()
    employer_date.paragraph_format.space_before = Pt(2)
    employer_date.paragraph_format.space_after = Pt(2)
    employer_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    employer_date_run = employer_date.add_run(f'{today.year}년      {today.month:02d}월      {today.day:02d}일')
    employer_date_run.font.size = Pt(8)
    employer_date_run.font.name = 'Malgun Gothic'
    
    employer_name = employer_info.get('대표자명', employer_info.get('회사명', '사업주'))
    employer_sign = doc.add_paragraph()
    employer_sign.paragraph_format.space_before = Pt(2)
    employer_sign.paragraph_format.space_after = Pt(6)
    employer_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    employer_sign_run = employer_sign.add_run(f'신청인(사업주)  {employer_name}  (서명 또는 인)')
    employer_sign_run.font.size = Pt(8)
    employer_sign_run.font.name = 'Malgun Gothic'
    
    # 제출처
    recipient = doc.add_paragraph()
    recipient.paragraph_format.space_before = Pt(2)
    recipient.paragraph_format.space_after = Pt(0)
    recipient_run = recipient.add_run('○○지방고용노동청(○○지청)장 귀하')
    recipient_run.font.size = Pt(8)
    recipient_run.font.name = 'Malgun Gothic'
    
    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# 1. 임신기 근로시간 단축 신청서 (한 페이지 완성)
# ============================================================

def create_application_form_docx(employee_info: Dict, employer_info: Dict,
                                  pregnancy_data: Dict, childbirth_data: Dict) -> BytesIO:
    """임신기 근로시간 단축 신청서 DOCX 생성 - 한 페이지 완성"""
    
    doc = Document()
    
    # 페이지 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run('임신기 근로시간 단축 신청서')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'Malgun Gothic'
    
    # 신청인 정보 테이블 (2행 5열)
    table1 = doc.add_table(rows=2, cols=5)
    table1.style = 'Table Grid'
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 첫 번째 열 병합 (신청인)
    cell_applicant = table1.cell(0, 0).merge(table1.cell(1, 0))
    cell_applicant.text = '신청인'
    cell_applicant.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_applicant.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell_applicant, "E7E6E6")
    
    # 헤더 셀 배경색
    shade_cell(table1.rows[0].cells[1], "F2F2F2")
    shade_cell(table1.rows[0].cells[3], "F2F2F2")
    shade_cell(table1.rows[1].cells[1], "F2F2F2")
    shade_cell(table1.rows[1].cells[3], "F2F2F2")
    
    # 첫째 행
    table1.rows[0].cells[1].text = '성명'
    table1.rows[0].cells[2].text = employee_info.get("이름", "")
    table1.rows[0].cells[3].text = '생년월일'
    # 주민등록번호 안전 처리
    birth_str = employee_info.get("주민등록번호", "")
    if birth_str and len(birth_str) >= 6:
        birth = birth_str[:6]
    else:
        birth = "000000"
    table1.rows[0].cells[4].text = birth
    
    # 둘째 행
    table1.rows[1].cells[1].text = '소속(부서)'
    table1.rows[1].cells[2].text = employee_info["부서"]
    table1.rows[1].cells[3].text = '직위(직급)'
    table1.rows[1].cells[4].text = employee_info["직급"]
    
    # 셀 서식 설정
    for row in table1.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Malgun Gothic'
    
    # 빈 줄
    spacer1 = doc.add_paragraph()
    spacer1.paragraph_format.space_before = Pt(8)
    spacer1.paragraph_format.space_after = Pt(8)
    
    # 임신기간 중 근로시간 단축 테이블
    due_date = childbirth_data["출산예정일"]
    start_date = pregnancy_data["시작일"]
    end_date = pregnancy_data["종료일"]
    work_start = pregnancy_data['근무시간']['시작']
    work_end = pregnancy_data['근무시간']['종료']
    
    # 복잡한 표 생성 (6행 4열)
    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 첫 번째 열: "임신기간 중 근로시간 단축" (모든 행 병합)
    cell_title = table2.cell(0, 0).merge(table2.cell(5, 0))
    cell_title.text = '임신기간 중\n근로시간\n단축'
    cell_title.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell_title, "E7E6E6")
    
    # 첫째 행: 출산예정일
    table2.cell(0, 1).merge(table2.cell(0, 2)).text = '출산예정일'
    table2.cell(0, 3).text = f"{due_date.year}년 {due_date.month:02d}월 {due_date.day:02d}일"
    shade_cell(table2.rows[0].cells[1], "F2F2F2")
    
    # 둘째-셋째 행: 12주(84일) 이내
    cell_12w = table2.cell(1, 1).merge(table2.cell(2, 1))
    cell_12w.text = '12주\n(84일)\n이내'
    shade_cell(cell_12w, "F2F2F2")
    
    shade_cell(table2.rows[1].cells[2], "F2F2F2")
    table2.rows[1].cells[2].text = '개시 예정일'
    table2.rows[1].cells[3].text = ''
    
    shade_cell(table2.rows[2].cells[2], "F2F2F2")
    table2.rows[2].cells[2].text = '종료 예정일'
    table2.rows[2].cells[3].text = ''
    
    # 넷째-다섯째 행: 32주(246일) 이후
    cell_32w = table2.cell(3, 1).merge(table2.cell(4, 1))
    cell_32w.text = '32주\n(246일)\n이후'
    shade_cell(cell_32w, "F2F2F2")
    
    shade_cell(table2.rows[3].cells[2], "F2F2F2")
    table2.rows[3].cells[2].text = '개시 예정일'
    table2.rows[3].cells[3].text = f"{start_date.year}.{start_date.month:02d}.{start_date.day:02d}"
    
    shade_cell(table2.rows[4].cells[2], "F2F2F2")
    table2.rows[4].cells[2].text = '종료 예정일'
    table2.rows[4].cells[3].text = f"{end_date.year}.{end_date.month:02d}.{end_date.day:02d}"
    
    # 여섯째 행: 근무 시간
    merged_cell = table2.cell(5, 1).merge(table2.cell(5, 2))
    merged_cell.text = '근무 개시 시각\n및 종료 시각'
    shade_cell(merged_cell, "F2F2F2")
    table2.cell(5, 3).text = f"{work_start} ~ {work_end}"
    
    # 셀 서식 설정
    for row in table2.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Malgun Gothic'
    
    # 주의사항
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(12)
    note_run = note.add_run('※ 개시 및 종료일정은 출산 일정에 따라 변동 될 수 있음.')
    note_run.font.size = Pt(9)
    note_run.font.name = 'Malgun Gothic'
    
    # 신청 문구
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(4)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1_run = p1.add_run('위 본인은 『근로기준법』제74조 제7항에 따라')
    p1_run.font.size = Pt(11)
    p1_run.font.name = 'Malgun Gothic'
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(12)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2_run = p2.add_run('위와 같이 근로시간 단축을 신청합니다.')
    p2_run.font.size = Pt(11)
    p2_run.font.name = 'Malgun Gothic'
    
    # 날짜
    today = datetime.now()
    date_p = doc.add_paragraph()
    date_p.paragraph_format.space_before = Pt(12)
    date_p.paragraph_format.space_after = Pt(4)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(f'{today.year}년      {today.month}월      {today.day}일')
    date_run.font.size = Pt(11)
    date_run.font.name = 'Malgun Gothic'
    
    # 신청인
    applicant_p = doc.add_paragraph()
    applicant_p.paragraph_format.space_before = Pt(0)
    applicant_p.paragraph_format.space_after = Pt(12)
    applicant_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    applicant_run = applicant_p.add_run(f"신청인  {employee_info['이름']}  (서명 또는 인)")
    applicant_run.font.size = Pt(11)
    applicant_run.font.name = 'Malgun Gothic'
    
    # 첨부
    attach_p = doc.add_paragraph()
    attach_p.paragraph_format.space_before = Pt(8)
    attach_p.paragraph_format.space_after = Pt(12)
    attach_run = attach_p.add_run('첨부  임신 사실을 증명하는 의사의 진단서(임신주수 확인용)')
    attach_run.font.size = Pt(9)
    attach_run.font.name = 'Malgun Gothic'
    
    # 제출처
    employer_name = employer_info.get('대표자명', employer_info.get('회사명', '사업주'))
    submit_p = doc.add_paragraph()
    submit_p.paragraph_format.space_before = Pt(8)
    submit_p.paragraph_format.space_after = Pt(0)
    submit_run = submit_p.add_run(f'{employer_name} 귀하')
    submit_run.font.size = Pt(10)
    submit_run.font.name = 'Malgun Gothic'
    
    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ============================================================
# 통합 함수
# ============================================================

def generate_pregnancy_forms_docx(employee_info: Dict, employer_info: Dict,
                                   pregnancy_data: Dict = None, childbirth_data: Dict = None) -> Dict[str, BytesIO]:
    """
    임신 관련 모든 서식을 DOCX로 생성
    """
    if pregnancy_data is None:
        pregnancy_data = C.PREGNANCY_SHORT_WORK
    
    if childbirth_data is None:
        childbirth_data = C.CHILDBIRTH_INFO
    
    forms = {}
    
    forms["임신기_근로시간_단축_신청서"] = create_application_form_docx(
        employee_info, employer_info, pregnancy_data, childbirth_data
    )
    
    forms["임신사유_근로시간_단축_확인서"] = create_confirmation_form_docx(
        employee_info, employer_info, pregnancy_data, childbirth_data
    )
    
    return forms
